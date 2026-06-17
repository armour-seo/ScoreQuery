"""
학술 논문 DOCX 생성 스크립트 (v2 — 교수자/학생 모드 개발 과정 포함)
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ── 페이지 설정 ──
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

# ── 스타일 ──
style = doc.styles['Normal']
font = style.font
font.name = '맑은 고딕'
font.size = Pt(10)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

IMG_DIR = r'C:\Users\armou\.gemini\antigravity\brain\73c42667-77af-4c2e-a0fc-367f4e420bc0'

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '맑은 고딕'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        if level == 1: run.font.size = Pt(14); run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
        elif level == 2: run.font.size = Pt(12); run.font.color.rgb = RGBColor(0x2d, 0x2d, 0x44)
        elif level == 3: run.font.size = Pt(11)
    return h

def add_para(text, bold=False, italic=False, size=None, align=None, space_after=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '맑은 고딕'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    if bold: run.bold = True
    if italic: run.italic = True
    if size: run.font.size = Pt(size)
    if align == 'center': p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right': p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if space_after is not None: p.paragraph_format.space_after = Pt(space_after)
    return p

def add_table(headers, rows, caption=None):
    if caption:
        add_para(caption, bold=True, size=9, space_after=4)
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True; run.font.size = Pt(8.5); run.font.name = '맑은 고딕'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        shading = cell._element.get_or_add_tcPr()
        shading.append(shading.makeelement(qn('w:shd'), {qn('w:val'):'clear', qn('w:color'):'auto', qn('w:fill'):'E8EAF6'}))
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8.5); run.font.name = '맑은 고딕'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    doc.add_paragraph()
    return table

def add_image(filename, width_inches=5.0, caption=None):
    path = os.path.join(IMG_DIR, filename) if not os.path.isabs(filename) else filename
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width_inches))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            add_para(caption, italic=True, size=9, align='center', space_after=8)

# ════════════════════════════════════════════════
# 문서 본문
# ════════════════════════════════════════════════

# ── 제목 ──
add_para('바이브 코딩을 이용한 교육용 시스템 개발 사례:', bold=True, size=15, align='center', space_after=2)
add_para('다국어 퀴즈 관리 시스템과 학생성적조회시스템', bold=True, size=15, align='center', space_after=12)

# ── Abstract ──
add_heading_styled('Abstract', level=2)
add_para(
    'The study reports two first-hand development cases in which an educator who felt uneasy about '
    '"coding" built the educational systems he needed by working with agentic AI through verification-centered '
    'vibe coding. UQM (Universal Quiz Master) is a high-complexity, long-pipeline system that integrates '
    'item authoring, AI item generation, instructor review, distribution, response collection, automatic scoring, '
    'analysis, and result release; it was developed over about three months. The Student Grade-Inquiry System '
    'is a short-pipeline tool that lets students review their own finalized results after grade processing; '
    'it was built within about one hour. Looking back on the two contrasting experiences, we find that the '
    'human role moves from writing code to defining problems, stating requirements, checking results, and '
    'giving final approval. We also find that the place where humans verify the work shifts with pipeline '
    'length: a single approval point in the short pipeline, and many spread-out checkpoints in the long one. '
    'We read vibe coding as a workable path to end-user development for non-IT educators, and we propose '
    'a verification-governance framework for using it responsibly.',
    italic=True, size=9, space_after=4
)
add_para('Keywords: agentic AI, vibe coding, end-user development, non-IT educators, verification governance',
         bold=True, size=9, space_after=12)

# ══════════════════════════════════
# 1. 서론
# ══════════════════════════════════
add_heading_styled('1. 서론', level=1)
add_para(
    '교수자는 날마다 디지털 도구가 필요하다. 수업을 운영하고, 평가를 치르고, 성적을 확인하고, 학생을 상담하고, '
    '학습 자료를 나누는 일에 모두 도구가 쓰인다. 정작 그 도구를 직접 만드는 일은 여전히 어렵게 느껴진다. '
    '특히 코딩이라는 말은 IT 전공이 아닌 교수자에게 큰 벽으로 다가온다.'
)
add_para(
    '이런 벽을 낮추는 변화가 최근 나타나고 있다. 생성형(Generative) AI와 주도형(Agentic) AI가 발전하면서, '
    '사용자는 문법을 먼저 배우지 않아도 원하는 기능을 말로 설명할 수 있게 되었다. AI가 만든 결과를 실행해 확인하고, '
    '잘못된 점은 다시 고쳐 달라고 요청한다. 이렇게 말로 풀어가는 개발 방식을 바이브 코딩(vibe coding)이라 부른다'
    '(Karpathy, 2025; Sarkar & Drosos, 2025). 바이브 코딩에서 중요한 능력은 코드를 처음부터 쓰는 솜씨가 아니다. '
    '풀고 싶은 문제를 또렷하게 설명하고, 결과를 확인하며, 잘못된 부분을 고칠 수 있는 요구로 바꾸는 능력이다.'
)
add_para(
    '이 논문은 주도형 AI가 교육용 시스템 개발의 문턱을 어떻게 낮추는지를 연구자 자신의 실제 개발 사례로 보여준다. '
    '다만 AI만 있으면 무엇이든 쉬워진다는 식의 장밋빛 주장은 하지 않는다. 서로 성격이 다른 두 사례를 나란히 놓고, '
    '균형 잡힌 그림을 그리고자 한다. 첫째 사례 UQM은 약 3개월이 걸린 길고 복잡한 통합 프로그램이고, 둘째 사례 '
    '학생성적조회시스템은 약 1시간 만에 완성한 짧고 실용적인 도구이다.'
)
add_para(
    '본 논문이 답하려는 물음은 세 가지이다. 첫째(Q1), 코딩을 잘 모르는 교수자도 주도형 AI를 써서 원하는 교육용 '
    '프로그램을 만들 수 있는가? 둘째(Q2), 길고 복잡한 프로그램과 짧고 단순한 시스템에서 개발 과정은 어떻게 달라지는가? '
    '셋째(Q3), 품질과 개인정보, 운영 책임을 지키기 위해 사람은 어떤 확인 역할을 해야 하는가?'
)

# ══════════════════════════════════
# 2. 이론적 배경
# ══════════════════════════════════
add_heading_styled('2. 이론적 배경', level=1)

add_heading_styled('2.1 생성형 AI, AI 에이전트, 주도형 AI', level=2)
add_para(
    '생성형 AI, AI 에이전트, 주도형 AI는 이어지는 발전 흐름 속에서 차례로 등장했고, 기능과 자율성에서 서로 '
    '구분된다(Sapkota et al., 2025). 생성형 AI는 사용자가 넣은 입력을 바탕으로 글, 그림, 코드, 문항 같은 새 '
    '결과물을 만든다. AI 에이전트는 도구를 써서 정해진 작업을 수행하는 실행 단위이다. 주도형 AI는 여러 에이전트가 '
    '협력하고, 큰 목표를 작은 일로 나누며, 목표에 따라 계획·실행·피드백·반복을 이어간다.'
)
add_table(
    ['구분', '생성형 AI', 'AI 에이전트', '주도형(Agentic) AI'],
    [
        ['핵심 기능', '입력→결과물 생성', '도구로 정해진 작업 수행', '에이전트 협력·일 나누기·기억 유지'],
        ['작동 방식', '요청→생성→출력', '지시→도구 사용→반환', '목표→계획→실행→피드백→반복'],
        ['사람의 역할', '요청하고 평가함', '지시하고 검토함', '목표를 정하고 확인·승인함'],
        ['주요 위험', '틀린 결과 생성', '작업 오류·도구 오작동', '책임 비움·지나친 위임'],
    ],
    caption='[표 1] 생성형 AI, AI 에이전트, 주도형 AI의 비교'
)

add_heading_styled('2.2 바이브 코딩과 검증 중심 개발', level=2)
add_para(
    '바이브 코딩은 사용자가 말로 기능을 설명하면 AI가 코드를 만들거나 고치고, 사용자가 그 결과를 확인한 뒤 '
    '다시 고쳐 달라고 요청하는 일을 되풀이하는 개발 방식이다(Sarkar & Drosos, 2025). 이 논문은 바이브 코딩을 '
    '검증 중심 바이브 코딩(verification-centered vibe coding)으로 새롭게 풀이한다.'
)
add_table(
    ['구분', '전통적 프로그래밍', '바이브 코딩'],
    [
        ['출발점', '문제 분석과 구조 설계', '원하는 기능과 결과 설명'],
        ['진행 구조', '설계→구현→테스트', '요청→생성→실행→확인→수정'],
        ['핵심 능력', '코드 작성·디버깅·구조 설계', '문제 정의·요구 표현·결과 확인'],
        ['AI의 역할', '보조 도구', '함께 만드는 개발 파트너'],
        ['주요 위험', '구현 오류·일정 지연', '결과를 그냥 믿음·확인 부족'],
    ],
    caption='[표 2] 전통적 프로그래밍과 바이브 코딩'
)

add_heading_styled('2.3 엔드유저 개발과 검증 거버넌스', level=2)
add_para(
    '엔드유저 개발(end-user development)은 전문 개발자가 아닌 최종 사용자가 자기 문제를 풀려고 소프트웨어를 '
    '직접 만들거나 고치는 일을 말한다(Ko et al., 2011; von Hippel, 2005). 이 논문은 확인과 승인의 책임 구조를 '
    '검증 거버넌스로 제안한다.'
)

# ══════════════════════════════════
# 3. 개발 사례 보고 방법
# ══════════════════════════════════
add_heading_styled('3. 개발 사례 보고 방법', level=1)
add_para(
    '이 논문은 연구자가 직접 만든 두 건의 교육용 시스템 개발 경험을 보고하고 돌아보는 개발 경험 보고이다. '
    '주도형 AI와 함께 시스템을 만들고 확인해 가는 사람-AI 협업 과정에 초점을 둔다.'
)
add_table(
    ['구분', 'UQM', '학생성적조회시스템'],
    [
        ['시스템 성격', '교육평가 운영 통합 프로그램', '확정 성적 기반 자기조회 시스템'],
        ['개발 환경', 'OpenAI Codex', 'Google Antigravity 2.0'],
        ['개발 기간', '약 3개월', '약 1시간'],
        ['작업 길이/복잡도', '긴 작업 / 매우 높음', '짧은 작업 / 낮음~중간'],
    ],
    caption='[표 3] 두 개발 사례 요약'
)

# ══════════════════════════════════
# 4. 개발 사례
# ══════════════════════════════════
add_heading_styled('4. 개발 사례', level=1)

# 4.1 UQM
add_heading_styled('4.1 개발 사례 1: UQM 다국어 퀴즈 관리 시스템', level=2)
add_para(
    'UQM은 교육평가 운영이 여기저기 흩어져 있고 손으로 하는 일이 많다는 문제에서 출발했다. 그동안은 출제, 배포, '
    '응답 수집, 채점, 분석, 결과 공개가 서로 다른 도구에 나뉘어 있었다. UQM은 이 일들을 하나로 묶으려고 연구자가 '
    '직접 만든 프로그램으로, 약 3개월이 걸린 어려운 작업이었다.'
)
add_table(
    ['구분', '도구/기술', '역할'],
    [
        ['개발 보조', 'ChatGPT, Codex', 'Python/Streamlit 코드 생성·오류 수정'],
        ['문항 생성', 'Gemini API', '객관식·OX·단답·서술형 문항 초안'],
        ['화면(UI)', 'Python, Streamlit', '교수자용 브라우저 화면'],
        ['문서 처리', 'pandas, openpyxl 등', 'PDF·PPT·엑셀·CSV 처리'],
        ['배포·수집', 'Google Forms API', '시험 배포와 응답 수집'],
    ],
    caption='[표 4] UQM의 주요 도구와 역할'
)
add_table(
    ['단계', '실제 사례'],
    [
        ['요구 정의', '[저자 기입: 통합 요구를 제시한 핵심 프롬프트]'],
        ['AI 생성', '[저자 기입: 만들어진 코드·구조 요약]'],
        ['실행 확인', '[저자 기입: 실행 결과]'],
        ['오류 수정', '[저자 기입: 대표적 오류와 수정 요청]'],
        ['검증', '[저자 기입: 문항 품질·번역·정답·배점·채점 확인]'],
        ['승인', '[저자 기입: 배포·공개를 승인한 근거]'],
    ],
    caption='[표 5] UQM 개발의 여섯 단계 사람-AI 상호작용 정리'
)

# ──────────────────────────────────
# 4.2 학생성적조회시스템 (교수자/학생 모드 전체)
# ──────────────────────────────────
add_heading_styled('4.2 개발 사례 2: 학생성적조회시스템', level=2)
add_para(
    '학생성적조회시스템은 성적을 확인하는 과정에서 거듭 생기는 불편을 풀려고 만들었다. 기능 범위가 또렷하고 입력 '
    '자료가 확정된 엑셀 파일로 한정되어, 약 1시간 만에 구현할 수 있었다. 이 시스템은 성적을 계산하는 도구가 아니다. '
    '교수자가 확정한 성적을 학생이 개인별로 안전하게 보는 피드백 도구이다. 시스템은 교수자 모드와 학생 모드의 '
    '두 갈래로 나뉘며, 각 모드의 역할과 기능이 분명히 구분된다.'
)

# 모드 선택 이미지
add_image('scorequery_mode_select_1781422576985.png', width_inches=4.5,
          caption='[그림 1] 시스템 진입 — 교수자/학생 모드 선택 화면')

# ── 4.2.1 교수자 모드 ──
add_heading_styled('4.2.1 교수자 모드: 설정에서 공시까지', level=3)
add_para(
    '교수자 모드는 성적 데이터를 준비하고 학생에게 공개하기까지의 전체 흐름을 4단계 마법사(wizard) 방식으로 안내한다. '
    '교수자는 코드를 다루지 않고 화면의 안내를 따라가면 된다.'
)

add_para('Step 1. 교수자 정보 입력', bold=True, size=10, space_after=2)
add_para(
    '첫 단계에서 교수자는 이름, 이메일, 연락처를 입력한다. 이 정보는 이후 학생 조회 화면의 상단 바에 '
    '"담당교수: 홍길동(hong@univ.ac.kr)" 형태로 표시되어, 학생이 문의처를 바로 알 수 있게 한다.'
)

add_para('Step 2. 과목 정보 등록', bold=True, size=10, space_after=2)
add_para(
    '교수자는 학년도, 학기, 과목명을 입력한다. 한 교수가 여러 과목을 담당하는 경우를 지원하기 위해 '
    '복수 과목 등록이 가능하다. 이미 등록된 과목명은 자동완성 목록으로 제시되어 재입력 부담을 줄인다. '
    '동일 과목(같은 년도·학기·과목명)을 다시 등록하면 "기존 설정을 덮어쓰시겠습니까?" 확인 대화상자가 나타나 '
    '실수로 기존 설정이 사라지는 것을 방지한다.'
)

add_para('Step 3. 평가 기준 설정', bold=True, size=10, space_after=2)
add_para(
    '퀴즈, 출석, 중간고사, 기말고사, 과제 등 평가 항목별 만점과 반영 비율을 설정한다. '
    '비율의 합이 100%가 되어야 다음 단계로 넘어갈 수 있다. 이 설정은 이후 엑셀 양식 자동 생성과 '
    '학생 화면의 점수 카드 구성에 그대로 반영된다.'
)

# 교수자 마법사 이미지
add_image('scorequery_admin_wizard_1781422536426.png', width_inches=4.5,
          caption='[그림 2] 교수자 모드 — 평가 기준 설정 (Step 3)')

add_para('Step 4. 완료 및 성적 데이터 관리', bold=True, size=10, space_after=2)
add_para(
    '마지막 단계에서는 과목별 설정 요약을 확인하고, 샘플 엑셀 양식을 다운로드할 수 있다. '
    '이 양식에는 앞서 설정한 평가 항목이 열(column)로 자동 반영되어, 교수자는 학생 정보와 점수만 채우면 된다. '
    '데이터베이스를 따로 두지 않고 엑셀을 그대로 데이터의 바탕으로 삼은 것은, 교수자가 이미 익숙한 도구 위에서 '
    '작업할 수 있도록 하기 위함이다.'
)

add_para('성적 데이터 업로드와 처리', bold=True, size=10, space_after=2)
add_para(
    '교수자는 완성된 엑셀 파일을 드래그앤드롭으로 업로드한다. 시스템은 업로드된 데이터를 자동으로 '
    '표준 JSON 형식으로 변환하며, 이 과정에서 학생의 이름은 부분 가림 처리(예: 홍*동), 학번은 일부 마스킹 '
    '처리되어 개인정보를 보호한다. 교수자는 [변환] → [미리보기] → [다운로드] → [확정]의 순서로 '
    '데이터를 단계적으로 검증한 뒤 확정한다.'
)

# 업로드 이미지
add_image('scorequery_admin_upload_1781422558603.png', width_inches=4.5,
          caption='[그림 3] 교수자 모드 — 성적 데이터 업로드 및 공시 설정')

add_para('성적 공시 기능', bold=True, size=10, space_after=2)
add_para(
    '교수자가 성적을 확정한 뒤에도 학생이 바로 볼 수 있는 것은 아니다. 과목별로 공시 일자와 시간을 지정해야 '
    '학생 조회가 열린다. 공시 전에 학생이 접속하면 "성적 공시 이전입니다"라는 안내 메시지가 표시된다. '
    '이 기능은 교수자가 최종 확인을 마친 뒤에만 성적이 노출되도록 하는 검증 거버넌스의 실현이다.'
)

# ── 4.2.2 학생 모드 ──
add_heading_styled('4.2.2 학생 모드: 인증에서 성적 확인까지', level=3)
add_para(
    '학생 모드는 세 단계로 진행된다. 과목 선택, 본인 인증, 성적 확인이다.'
)

add_para('과목 선택과 인증', bold=True, size=10, space_after=2)
add_para(
    '학생은 먼저 학년도, 학기, 과목을 선택한다. 미선택 상태의 드롭다운은 옅은 노란색 배경으로 표시되어 '
    '선택이 필요함을 시각적으로 안내한다. 과목을 선택하면 학번(ID)과 전화번호 뒷자리 4자리를 입력하는 '
    '인증 영역이 나타난다. 인증은 교수자가 업로드한 엑셀의 학생 정보와 대조하여 이루어진다. '
    '인증 정보는 서버에 저장되지 않으며, 본인 확인 목적으로만 사용된다.'
)

# 로그인 이미지
add_image('scorequery_login_1781422088944.png', width_inches=3.5,
          caption='[그림 4] 학생 모드 — 과목 선택 및 로그인 화면')

add_para('성적 결과 화면', bold=True, size=10, space_after=2)
add_para(
    '인증에 성공하면 상단 바가 "성적조회시스템: 2026-1학기-경영정보시스템 | 담당교수: 홍길동(email)"으로 '
    '바뀌고, 성적 결과 화면이 나타난다. 화면은 다음과 같이 구성된다.'
)
add_para(
    '첫째, 학생 정보 헤더에는 이름(가림 처리), 학과, 분반, 학번(마스킹)이 표시된다. '
    '둘째, 평가 항목별 점수 카드가 나열되며, 각 카드에는 본인 점수, 만점 기준, 프로그레스 바, 분반 평균, '
    '그리고 분반 평균과의 차이가 +녹색/-노란색 볼드로 표시되어 한눈에 위치를 파악할 수 있다. '
    '셋째, 레이더 차트가 내 점수·분반 평균·분반 최고점을 겹쳐 보여주어 강점과 약점을 시각적으로 비교할 수 있다. '
    '넷째, 종합 정보에는 평점, 분반 석차, 총점, 결석 횟수가 요약되고, 하단에는 성적 이의신청 안내가 표시된다.'
)

# 결과 이미지
add_image('scorequery_result_1781422070571.png', width_inches=5.0,
          caption='[그림 5] 학생 모드 — 성적 결과 화면 (점수 카드, 레이더 차트, 종합 정보)')

# ── 4.2.3 시스템 구조 ──
add_heading_styled('4.2.3 시스템 구조와 기술 스택', level=3)
add_table(
    ['구분', '기술/도구', '역할'],
    [
        ['프론트엔드', 'HTML, CSS, JavaScript', '교수자/학생 화면, 인증, 차트'],
        ['차트', 'Chart.js', '레이더 차트 (반 평균·최고점 비교)'],
        ['엑셀 처리', 'SheetJS (xlsx)', '엑셀 업로드 파싱, 샘플 양식 생성'],
        ['데이터 저장', 'localStorage + JSON', '과목 설정·성적 데이터 저장'],
        ['개인정보 보호', '해시 기반 인증', '전화번호 뒷자리 SHA-256 해시 비교'],
        ['배포', 'GitHub Pages', '정적 호스팅, 서버 불필요'],
        ['개발 도구', 'Google Antigravity 2.0', '바이브 코딩 기반 전체 개발'],
    ],
    caption='[표 6] 학생성적조회시스템의 기술 스택'
)

# ── 4.2.4 개발 과정의 여섯 단계 ──
add_heading_styled('4.2.4 개발 과정의 여섯 단계 정리', level=3)
add_para(
    '다음 표는 학생성적조회시스템의 개발 과정을 요구 정의에서 승인까지 여섯 단계로 정리한 것이다. '
    '교수자 모드와 학생 모드의 주요 기능이 어떤 요구에서 출발하여 어떤 확인을 거쳐 완성되었는지를 보여준다.'
)
add_table(
    ['단계', '교수자 모드', '학생 모드'],
    [
        ['요구 정의',
         '과목 설정 마법사, 엑셀 업로드,\n공시 일시 지정 요구',
         '과목 선택 후 학번+전화번호로\n본인 성적만 조회 요구'],
        ['AI 생성',
         '4단계 위자드 UI, 드래그앤드롭\n업로드, 변환·미리보기·확정 파이프라인',
         '과목 드롭다운, 인증 로직,\n점수 카드, 레이더 차트, 종합 정보'],
        ['실행 확인',
         '과목 등록·평가비율 100% 검증,\n엑셀→JSON 변환 결과 확인',
         '로그인 동작, 카드 표시,\n차트 렌더링 확인'],
        ['오류 수정',
         '동일 과목 덮어쓰기 시 경고 누락,\n과목 추가 시 이전 과목 잔존',
         'localStorage 오래된 데이터가\n정상 데이터를 차단하는 문제'],
        ['검증',
         '엑셀 원본과 JSON 일치 확인,\n가림 처리 적정성, 공시 동작',
         '점수·석차·평균 원본 일치,\n본인 외 정보 노출 없음 확인'],
        ['승인',
         '교수자가 [확정] 후 공시 일시 설정',
         '공시 이후에만 학생 조회 허용'],
    ],
    caption='[표 7] 학생성적조회시스템 개발의 여섯 단계 정리 (교수자/학생 모드)'
)

# ── 4.3 대조 ──
add_heading_styled('4.3 두 개발 사례의 대조', level=2)
add_para(
    '바이브 코딩은 모든 개발을 쉽게 만들지 않는다. 복잡한 프로그램은 여전히 어렵다. 그렇더라도 주도형 AI는 비IT '
    '사용자가 그 개발에 들어설 길을 열어준다.'
)
add_table(
    ['구분', 'UQM', '학생성적조회시스템'],
    [
        ['개발 기간', '약 3개월', '약 1시간'],
        ['작업 흐름', '준비-출제-배포-수집-채점-분석-공개', '과목 설정-업로드-공시-조회'],
        ['복잡도', '매우 높음', '낮음~중간'],
        ['AI에 맡긴 범위', '코드·문항·데이터·배포·분석', '교수자 위자드·학생 조회·차트'],
        ['사람 확인 지점', '문항·번역·배점·채점·공개(여러 곳)', '원본 일치·인증·공시·개인정보'],
        ['경험의 뜻', '길고 복잡한 작업에도 쓸 수 있음', '빠르게 실용화할 수 있음'],
    ],
    caption='[표 8] 두 개발 사례의 대조'
)

# ══════════════════════════════════
# 5. 결론
# ══════════════════════════════════
add_heading_styled('5. 결론', level=1)

add_heading_styled('5.1 개발 사례의 요약과 시사점', level=2)
add_para(
    '이 논문은 코딩에 부담을 느끼던 교수자에게 주도형 AI 기반 바이브 코딩이 새로운 길이 될 수 있음을 두 개발 '
    '사례로 보고했다. 핵심은 주도형 AI가 개발을 무조건 쉽게 만든다는 것이 아니다. 엄두를 내지 못하던 교수자에게 '
    '일단 시작할 길을 열어준다는 것이다.'
)
add_para(
    '두 개발 경험을 바탕으로 세 가지 명제를 제안한다. 명제 1: 작업이 길어질수록 사람의 확인은 한 곳의 승인에서 '
    '여러 곳으로 흩어진다. 명제 2: 복잡도가 높아질수록 도메인 지식을 갖춘 사람의 개입이 성공을 가른다. '
    '명제 3: 확인 기준에 교육 고유의 책임이 들어갈수록 엔드유저의 승인 책임이 시스템 신뢰를 좌우한다.'
)

add_heading_styled('5.2 검증 거버넌스 프레임워크', level=2)
add_para(
    '검증 거버넌스는 AI를 책임 있게 쓰기 위한 조건이다. AI가 낸 결과를 그냥 믿지 않고, 원본과 견주고, '
    '개인정보를 점검하고, 교육적으로 타당한지 확인할 때 AI는 믿을 수 있는 개발 파트너가 된다.'
)
add_table(
    ['차원', '핵심 질문', 'UQM(긴 작업)', '학생성적조회(짧은 작업)'],
    [
        ['요구사항', '무엇을 만들 것인가?', '출제~공개 전 과정 정의', '교수자 설정+학생 자기조회'],
        ['AI에 맡김', '어디까지 맡길 것인가?', '코드·문항·데이터·분석', '위자드·조회 로직·차트'],
        ['데이터 확인', '원본과 맞는가?', '문항·정답키·응답·채점표', '엑셀→JSON 변환 결과 일치'],
        ['품질 확인', '교육적으로 타당한가?', '난이도·번역·배점·채점 오류', '점수 표시·비교·가독성'],
        ['개인정보', '정보가 드러나는가?', '응답·계정·API 키 보호', '가림 처리·해시 인증·공시 제한'],
        ['승인 책임', '누가 판단하는가?', '교수자 승인 뒤 배포', '교수자 공시 뒤 조회 허용'],
        ['유지보수', '오류 시 어떻게?', '버전 관리·단계별 추적', '엑셀 수정→재업로드→재확인'],
    ],
    caption='[표 9] 검증 거버넌스 프레임워크'
)

add_heading_styled('5.3 한계와 향후 과제', level=2)
add_para(
    '이 논문은 두 건의 개발 사례에 기댄 보고이자 연구자 자신의 경험이다. 결과를 모든 현장으로 넓히기는 어렵고, '
    '해석이 치우칠 여지도 있다. 앞으로는 여러 교수자의 개발 사례를 모으고, 제3자 검토를 받아 신뢰를 높일 필요가 있다. '
    '개발 시간, 되풀이 횟수, 오류율, 만족도를 숫자로 분석하는 일도 남는다. 만들어진 코드의 보안 약점, 유지보수, '
    '접근성도 전문가가 따로 살펴야 한다.'
)

# GitHub 배포 이미지
add_image('media__1781421510549.png', width_inches=5.0,
          caption='[그림 6] GitHub Pages를 통한 배포 — 레포지토리 화면')

# ── 저장 ──
output_path = r'c:\ScoreQuery\바이브코딩_교육용시스템_개발사례.docx'
doc.save(output_path)
print(f'DOCX 저장 완료: {output_path}')
